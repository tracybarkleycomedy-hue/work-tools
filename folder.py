import io
import re
import time
import requests
import streamlit as st
from docx import Document
from html import unescape


st.set_page_config(page_title="Client Folder + Personality Generator", layout="wide")


# --------------------------------------------------
# Basic helpers
# --------------------------------------------------
def clean_folder_name(text):
    text = text.strip()
    text = re.sub(r'[<>:"/\\|?*]', "", text)
    text = re.sub(r"\s+", " ", text)
    return text


def normalize_site_path(site_path):
    site_path = site_path.strip()
    site_path = site_path.replace("https://", "").replace("http://", "")
    if "/sites/" in site_path:
        site_path = site_path.split("/sites/", 1)[1]
    return site_path.strip("/")


def replace_placeholders(text, client_name, client_id):
    if not text:
        return text

    replacements = {
        "{{CLIENT_NAME}}": client_name,
        "{{CLIENT_ID}}": client_id,
        "{CLIENT_NAME}": client_name,
        "{CLIENT_ID}": client_id,
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    return text


# --------------------------------------------------
# Microsoft Graph helpers
# --------------------------------------------------
def get_graph_token(tenant_id, client_id, client_secret):
    url = f"https://login.microsoftonline.com/{tenant_id}/oauth2/v2.0/token"

    data = {
        "client_id": client_id,
        "client_secret": client_secret,
        "scope": "https://graph.microsoft.com/.default",
        "grant_type": "client_credentials",
    }

    response = requests.post(url, data=data, timeout=20)

    if response.status_code >= 400:
        raise Exception(f"Token error {response.status_code}: {response.text}")

    return response.json()["access_token"]


def graph_get(url, token):
    response = requests.get(
        url,
        headers={"Authorization": f"Bearer {token}"},
        timeout=30,
    )

    if response.status_code >= 400:
        raise Exception(f"Graph API error {response.status_code} for GET\n{url}\n{response.text}")

    return response.json()


def graph_post(url, token, body):
    response = requests.post(
        url,
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json",
        },
        json=body,
        timeout=30,
    )

    if response.status_code >= 400:
        raise Exception(f"Graph API error {response.status_code} for POST\n{url}\n{response.text}")

    return response


def graph_patch(url, token, body):
    response = requests.patch(
        url,
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json",
        },
        json=body,
        timeout=30,
    )

    if response.status_code >= 400:
        raise Exception(f"Graph API error {response.status_code} for PATCH\n{url}\n{response.text}")

    return response.json() if response.text else {}


def graph_put_content(url, token, content_bytes):
    response = requests.put(
        url,
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        },
        data=content_bytes,
        timeout=60,
    )

    if response.status_code >= 400:
        raise Exception(f"Graph API error {response.status_code} for PUT\n{url}\n{response.text}")

    return response.json() if response.text else {}


def graph_download(url, token):
    response = requests.get(
        url,
        headers={"Authorization": f"Bearer {token}"},
        timeout=60,
    )

    if response.status_code >= 400:
        raise Exception(f"Graph API error {response.status_code} for DOWNLOAD\n{url}\n{response.text}")

    return response.content


def get_site_id(hostname, site_path, token):
    site_path = normalize_site_path(site_path)
    url = f"https://graph.microsoft.com/v1.0/sites/{hostname}:/sites/{site_path}"
    data = graph_get(url, token)
    return data["id"]


def get_drive_id(site_id, drive_name, token):
    url = f"https://graph.microsoft.com/v1.0/sites/{site_id}/drives"
    data = graph_get(url, token)

    for drive in data.get("value", []):
        if drive.get("name", "").lower() == drive_name.lower():
            return drive["id"]

    available = [d.get("name") for d in data.get("value", [])]
    raise Exception(f"Drive '{drive_name}' not found. Available drives: {available}")


def get_item_by_path(drive_id, path, token):
    path = path.strip("/")
    url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/root:/{path}"
    return graph_get(url, token)


def list_children(drive_id, item_id, token):
    url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}/children"
    data = graph_get(url, token)
    return data.get("value", [])


def wait_for_copied_folder(drive_id, parent_item_id, new_folder_name, token, max_wait_seconds=90):
    start = time.time()

    while time.time() - start < max_wait_seconds:
        children = list_children(drive_id, parent_item_id, token)

        for child in children:
            if child.get("name") == new_folder_name:
                return child

        time.sleep(5)

    raise Exception("Folder copy was submitted, but the new folder was not found before timeout. SharePoint may still be copying it.")


def rename_item(drive_id, item_id, new_name, token):
    url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item_id}"
    return graph_patch(url, token, {"name": new_name})


def update_docx_placeholders(drive_id, item, token, client_name, client_id):
    download_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item['id']}/content"
    content = graph_download(download_url, token)

    doc_stream = io.BytesIO(content)
    doc = Document(doc_stream)

    def replace_in_paragraph(paragraph):
        full_text = "".join(run.text for run in paragraph.runs)
        new_text = replace_placeholders(full_text, client_name, client_id)

        if new_text != full_text:
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    replace_in_paragraph(paragraph)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    upload_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{item['id']}/content"
    graph_put_content(upload_url, token, output.getvalue())


def customize_copied_folder(drive_id, folder_item, token, client_name, client_id, log):
    children = list_children(drive_id, folder_item["id"], token)

    for child in children:
        original_name = child.get("name", "")
        new_name = replace_placeholders(original_name, client_name, client_id)

        if new_name != original_name:
            child = rename_item(drive_id, child["id"], new_name, token)
            log.append(f"Renamed: {original_name} → {new_name}")

        is_folder = "folder" in child
        is_docx = child.get("name", "").lower().endswith(".docx")

        if is_docx:
            update_docx_placeholders(drive_id, child, token, client_name, client_id)
            log.append(f"Updated Word doc placeholders: {child.get('name')}")

        if is_folder:
            customize_copied_folder(drive_id, child, token, client_name, client_id, log)


def create_folder_from_template(
    tenant_id,
    client_id,
    client_secret,
    hostname,
    site_path,
    drive_name,
    parent_folder_path,
    template_folder_name,
    new_folder_name,
    real_client_name,
    real_client_id,
):
    log = []

    token = get_graph_token(tenant_id, client_id, client_secret)
    log.append("Authenticated with Microsoft Graph.")

    site_id = get_site_id(hostname, site_path, token)
    log.append("Connected to SharePoint site.")

    drive_id = get_drive_id(site_id, drive_name, token)
    log.append(f"Connected to document library: {drive_name}")

    parent_folder_path = parent_folder_path.strip("/")
    template_path = f"{parent_folder_path}/{template_folder_name}".strip("/")

    template_item = get_item_by_path(drive_id, template_path, token)
    parent_item = get_item_by_path(drive_id, parent_folder_path, token)

    copy_url = f"https://graph.microsoft.com/v1.0/drives/{drive_id}/items/{template_item['id']}/copy"

    body = {
        "parentReference": {
            "driveId": drive_id,
            "id": parent_item["id"],
        },
        "name": new_folder_name,
    }

    graph_post(copy_url, token, body)
    log.append("Folder copy request submitted.")

    copied_folder = wait_for_copied_folder(drive_id, parent_item["id"], new_folder_name, token)
    log.append("Copied folder found. Starting customization.")

    customize_copied_folder(
        drive_id=drive_id,
        folder_item=copied_folder,
        token=token,
        client_name=real_client_name,
        client_id=real_client_id,
        log=log,
    )

    log.append("Customization complete.")

    return {
        "message": "Folder copied and customized successfully.",
        "new_folder_name": new_folder_name,
        "log": log,
    }


# --------------------------------------------------
# Website / personality helpers
# --------------------------------------------------
def fetch_website_text(url):
    if not url:
        return ""

    site = url.strip()
    if not site.startswith("http://") and not site.startswith("https://"):
        site = f"https://{site}"

    try:
        response = requests.get(
            site,
            timeout=12,
            headers={"User-Agent": "Mozilla/5.0"},
        )
        html = response.text

        title_match = re.search(r"<title[^>]*>(.*?)</title>", html, re.I | re.S)
        title = unescape(title_match.group(1)).strip() if title_match else ""

        meta_match = re.search(
            r'<meta[^>]+name=["\']description["\'][^>]+content=["\'](.*?)["\']',
            html,
            re.I | re.S,
        )
        meta = unescape(meta_match.group(1)).strip() if meta_match else ""

        text = re.sub(r"<script.*?</script>", " ", html, flags=re.I | re.S)
        text = re.sub(r"<style.*?</style>", " ", text, flags=re.I | re.S)
        text = re.sub(r"<[^>]+>", " ", text)
        text = unescape(text)
        text = re.sub(r"\s+", " ", text).strip()

        combined = f"{title}. {meta}. {text}"
        return combined[:4000]

    except Exception:
        return ""


def guess_focus_words(text):
    lower = text.lower()
    focus = []

    keyword_map = {
        "association": "members and industry professionals",
        "member": "members and stakeholders",
        "education": "learners, educators, and professionals",
        "health": "patients, professionals, and care teams",
        "medical": "patients, clinicians, and healthcare professionals",
        "conference": "attendees, exhibitors, and members",
        "certification": "professionals pursuing certification and education",
        "advocacy": "members, policymakers, and the public",
        "research": "researchers, professionals, and informed readers",
        "community": "members and community stakeholders",
    }

    for key, phrase in keyword_map.items():
        if key in lower and phrase not in focus:
            focus.append(phrase)

    if not focus:
        focus.append("members, customers, and stakeholders")

    return focus[0]


def generate_personality_prompt(
    assistant_name,
    org_name,
    acronym,
    website,
    industry,
    audience,
    tone_words,
    mission,
    topics,
    extra_notes,
    website_text,
):
    assistant_name = assistant_name.strip() or "Betty"
    org_name = org_name.strip() or "the organization"
    acronym_text = f" ({acronym.strip()})" if acronym.strip() else ""
    industry = industry.strip() or "the organization’s field"
    audience = audience.strip() or guess_focus_words(website_text)
    tone_words = tone_words.strip() or "warm, credible, and helpful"
    mission = mission.strip() or "supporting its members, customers, and broader community"
    topics = topics.strip() or f"{industry}, member needs, available resources, and organizational services"

    if extra_notes.strip():
        extra = f" You also keep in mind: {extra_notes.strip()}"
    else:
        extra = ""

    prompt = (
        f"Hey {assistant_name}! You’re the {tone_words} voice of {org_name}{acronym_text} — "
        f"passionate about {mission}. You speak with clarity about {topics}. "
        f"You balance accurate, source-informed guidance with an inviting tone that makes complex concepts feel approachable and useful. "
        f"Your goal is to help {audience} find trusted information, understand available resources, and feel confident taking the next step with {org_name}.{extra}"
    )

    return prompt


def generate_greeting_prompt(assistant_name, org_name, audience):
    assistant_name = assistant_name.strip() or "Betty"
    org_name = org_name.strip() or "the organization"
    audience = audience.strip() or "members, customers, and visitors"

    return (
        f"Hi, I’m {assistant_name}, your AI assistant for {org_name}. "
        f"I can help {audience} find information, understand available resources, and get pointed in the right direction. "
        f"What can I help you with today?"
    )


# --------------------------------------------------
# UI
# --------------------------------------------------
st.title("📁 Client Folder Tool + AI Personality Generator")
st.caption("Create client onboarding folders, customize placeholders, and generate Betty personality prompts.")

tab1, tab2 = st.tabs(["📁 Folder Creator", "🤖 Personality Prompt Generator"])


# --------------------------------------------------
# Folder creator tab
# --------------------------------------------------
with tab1:
    st.header("Client Folder Creator")
    st.write("Create and customize a new client folder from your onboarding template.")

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("Client Info")
        client_name = st.text_input("Client Name")
        client_number = st.text_input("Client ID / Number")

        folder_format = st.selectbox(
            "Folder Name Format",
            [
                "Client Name - Client ID",
                "Client ID - Client Name",
                "Client Name Only",
                "Custom",
            ],
        )

        custom_folder_name = ""
        if folder_format == "Custom":
            custom_folder_name = st.text_input("Custom Folder Name")

    with col2:
        st.subheader("Connection Settings")
        tenant_id = st.text_input("Tenant ID")
        app_client_id = st.text_input("App Client ID")
        client_secret = st.text_input("Client Secret", type="password")

        sharepoint_hostname = st.text_input("SharePoint Hostname", placeholder="bluecypress01.sharepoint.com")
        site_path = st.text_input("Site Path", placeholder="BlueCypress-ChatGPT")
        drive_name = st.text_input("Drive Name", value="Documents")
        parent_folder_path = st.text_input("Parent Folder Path", value="Client Folders/_ONBOARDING")
        template_folder_name = st.text_input("Template Folder Name", value="1 Folder Template-New Clients")

    if folder_format == "Client Name - Client ID":
        new_folder_name = clean_folder_name(f"{client_name} - {client_number}")
    elif folder_format == "Client ID - Client Name":
        new_folder_name = clean_folder_name(f"{client_number} - {client_name}")
    elif folder_format == "Client Name Only":
        new_folder_name = clean_folder_name(client_name)
    else:
        new_folder_name = clean_folder_name(custom_folder_name)

    st.info(f"New folder will be named: **{new_folder_name or '[enter client info]'}**")

    st.caption("This will copy the template folder, then rename child folders/files and replace {{CLIENT_NAME}} and {{CLIENT_ID}} inside Word docs.")

    if st.button("Create Folder", type="primary"):
        if not new_folder_name:
            st.error("Please enter enough client information to create a folder name.")
        elif not client_name:
            st.error("Please enter Client Name.")
        elif not client_number:
            st.error("Please enter Client ID / Number.")
        else:
            try:
                with st.spinner("Copying template and customizing placeholders..."):
                    result = create_folder_from_template(
                        tenant_id=tenant_id,
                        client_id=app_client_id,
                        client_secret=client_secret,
                        hostname=sharepoint_hostname,
                        site_path=site_path,
                        drive_name=drive_name,
                        parent_folder_path=parent_folder_path,
                        template_folder_name=template_folder_name,
                        new_folder_name=new_folder_name,
                        real_client_name=client_name,
                        real_client_id=client_number,
                    )

                st.success(result["message"])
                st.write(f"Created/customized folder: **{result['new_folder_name']}**")

                with st.expander("Process log"):
                    for item in result["log"]:
                        st.write(f"- {item}")

            except Exception as e:
                st.error(str(e))


# --------------------------------------------------
# Personality generator tab
# --------------------------------------------------
with tab2:
    st.header("AI Personality Prompt Generator")
    st.write("Generate a ready-to-paste Betty personality block from a website and client details.")

    pcol1, pcol2 = st.columns(2)

    with pcol1:
        assistant_name = st.text_input("Assistant Name", value="Betty")
        org_name = st.text_input("Organization Name")
        acronym = st.text_input("Acronym / Short Name", placeholder="AANP")
        website = st.text_input("Website", placeholder="https://example.org")
        industry = st.text_input("Industry / Field", placeholder="healthcare, association, finance, education...")

    with pcol2:
        audience = st.text_input("Primary Audience", placeholder="members, patients, professionals, customers...")
        tone_words = st.text_input("Tone Words", value="warm, credible, and helpful")
        mission = st.text_area("Mission / Purpose", placeholder="empowering members, advancing the profession, helping customers...", height=90)
        topics = st.text_area("Topics Betty Should Speak Clearly About", placeholder="member resources, services, events, education, benefits...", height=90)
        extra_notes = st.text_area("Extra Notes", placeholder="Anything specific about voice, audience, or client expectations...", height=90)

    if st.button("Analyze Website + Generate Personality", type="primary"):
        website_text = fetch_website_text(website)

        personality = generate_personality_prompt(
            assistant_name=assistant_name,
            org_name=org_name,
            acronym=acronym,
            website=website,
            industry=industry,
            audience=audience,
            tone_words=tone_words,
            mission=mission,
            topics=topics,
            extra_notes=extra_notes,
            website_text=website_text,
        )

        greeting = generate_greeting_prompt(
            assistant_name=assistant_name,
            org_name=org_name,
            audience=audience,
        )

        st.divider()

        st.subheader("Personality Prompt")
        st.code(personality)

        st.subheader("Greeting Prompt")
        st.code(greeting)

        st.subheader("Suggested Placement")
        st.write("Personality prompt: **PERSONALITY → Instructions**")
        st.write("Greeting prompt: **GREETING → Instructions**")

        st.subheader("Website Scan Notes")
        if website_text:
            st.caption("Website text was detected and used lightly for audience/context hints.")
            with st.expander("Preview website text used"):
                st.write(website_text[:1500])
        else:
            st.info("No website text could be pulled. The prompt was generated from the fields you entered.")